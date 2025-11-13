"""
Serwis generowania raportów DOCX (Microsoft Word).
Używa python-docx do tworzenia dokumentów .docx.
"""
import io
import logging
from datetime import datetime
from typing import Dict, Any, List

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger(__name__)


class DOCXGenerator:
    """Generator raportów DOCX dla person, grup fokusowych i ankiet."""

    def __init__(self):
        """Inicjalizacja generatora."""
        pass

    async def generate_persona_docx(
        self,
        persona_data: Dict[str, Any],
        user_tier: str = "free",
        include_reasoning: bool = True,
    ) -> bytes:
        """
        Generuje raport DOCX dla persony.

        Args:
            persona_data: Dane persony
            user_tier: Tier użytkownika ('free', 'pro', 'enterprise')
            include_reasoning: Czy dołączyć reasoning

        Returns:
            bytes: Zawartość pliku DOCX
        """
        logger.info(f"Generowanie DOCX dla persony: {persona_data.get('name', 'Unknown')}")

        doc = Document()

        # Watermark dla free tier
        if user_tier == "free":
            self._add_watermark(doc, "SIGHT FREE TIER")

        # Tytuł
        title = doc.add_heading(f"Raport Persony: {persona_data.get('name', 'Unknown')}", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Metadata
        metadata = doc.add_paragraph()
        metadata.add_run(f"Wygenerowano: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n").italic = True
        metadata.add_run("Platform: Sight AI Focus Groups").italic = True

        # Dane podstawowe
        doc.add_heading("Dane Podstawowe", level=2)
        table = doc.add_table(rows=6, cols=2)
        table.style = "Light Grid Accent 1"

        fields = [
            ("Imię", persona_data.get("name", "N/A")),
            ("Wiek", str(persona_data.get("age", "N/A"))),
            ("Zawód", persona_data.get("occupation", "N/A")),
            ("Wykształcenie", persona_data.get("education", "N/A")),
            ("Dochód", persona_data.get("income", "N/A")),
            ("Lokalizacja", persona_data.get("location", "N/A")),
        ]

        for i, (label, value) in enumerate(fields):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)

        # Profil psychologiczny
        doc.add_heading("Profil Psychologiczny", level=2)

        if persona_data.get("values"):
            p = doc.add_paragraph()
            p.add_run("Wartości: ").bold = True
            p.add_run(", ".join(persona_data["values"]))

        if persona_data.get("interests"):
            p = doc.add_paragraph()
            p.add_run("Zainteresowania: ").bold = True
            p.add_run(", ".join(persona_data["interests"]))

        # Potrzeby
        if persona_data.get("needs"):
            doc.add_heading("Potrzeby", level=2)
            for need in persona_data["needs"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{need.get('category', 'N/A')}: ").bold = True
                p.add_run(f"{need.get('description', 'N/A')} ")
                p.add_run(f"(priorytet: {need.get('priority', 'N/A')})").italic = True

        # Reasoning (jeśli include_reasoning)
        if include_reasoning and persona_data.get("reasoning"):
            doc.add_heading("Uzasadnienie Generacji", level=2)
            reasoning = persona_data["reasoning"]

            if reasoning.get("orchestration_brief"):
                doc.add_paragraph(reasoning["orchestration_brief"])

            if reasoning.get("graph_insights"):
                doc.add_heading("Insighty z Graph RAG", level=3)
                doc.add_paragraph(reasoning["graph_insights"])

        # Konwertuj do bytes
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_bytes = docx_buffer.getvalue()

        logger.info(f"DOCX wygenerowany, rozmiar: {len(docx_bytes)} bytes")
        return docx_bytes

    async def generate_focus_group_docx(
        self,
        focus_group_data: Dict[str, Any],
        user_tier: str = "free",
        include_discussion: bool = True,
    ) -> bytes:
        """
        Generuje raport DOCX dla grupy fokusowej.

        Args:
            focus_group_data: Dane grupy fokusowej
            user_tier: Tier użytkownika
            include_discussion: Czy dołączyć transkrypcję

        Returns:
            bytes: Zawartość pliku DOCX
        """
        logger.info(f"Generowanie DOCX dla grupy fokusowej: {focus_group_data.get('name', 'Unknown')}")

        doc = Document()

        if user_tier == "free":
            self._add_watermark(doc, "SIGHT FREE TIER")

        # Tytuł
        title = doc.add_heading(f"Raport Grupy Fokusowej: {focus_group_data.get('name', 'Unknown')}", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Metadata
        metadata = doc.add_paragraph()
        metadata.add_run(f"Wygenerowano: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n").italic = True
        metadata.add_run("Platform: Sight AI Focus Groups").italic = True

        # Informacje podstawowe
        doc.add_heading("Informacje Podstawowe", level=2)
        p = doc.add_paragraph()
        p.add_run("Nazwa: ").bold = True
        p.add_run(focus_group_data.get("name", "N/A"))

        personas_count = len(focus_group_data.get("personas", []))
        p = doc.add_paragraph()
        p.add_run("Liczba uczestników: ").bold = True
        p.add_run(str(personas_count))

        # Podsumowanie
        if focus_group_data.get("summary"):
            doc.add_heading("Podsumowanie", level=2)
            summary = focus_group_data["summary"]

            if summary.get("main_insights"):
                doc.add_paragraph(summary["main_insights"])

            if summary.get("key_themes"):
                doc.add_heading("Kluczowe Tematy", level=3)
                for theme in summary["key_themes"]:
                    doc.add_paragraph(theme, style="List Bullet")

        # Transkrypcja dyskusji
        if include_discussion and focus_group_data.get("discussion"):
            doc.add_heading("Transkrypcja Dyskusji", level=2)

            for message in focus_group_data["discussion"]:
                p = doc.add_paragraph()
                p.add_run(f"{message.get('persona_name', 'Unknown')}: ").bold = True
                p.add_run(message.get("content", ""))

        # Konwertuj do bytes
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_bytes = docx_buffer.getvalue()

        logger.info(f"DOCX wygenerowany, rozmiar: {len(docx_bytes)} bytes")
        return docx_bytes

    async def generate_survey_docx(
        self,
        survey_data: Dict[str, Any],
        user_tier: str = "free",
    ) -> bytes:
        """
        Generuje raport DOCX dla ankiety.

        Args:
            survey_data: Dane ankiety z odpowiedziami
            user_tier: Tier użytkownika

        Returns:
            bytes: Zawartość pliku DOCX
        """
        logger.info(f"Generowanie DOCX dla ankiety: {survey_data.get('title', 'Unknown')}")

        doc = Document()

        if user_tier == "free":
            self._add_watermark(doc, "SIGHT FREE TIER")

        # Tytuł
        title = doc.add_heading(f"Raport Ankiety: {survey_data.get('title', 'Unknown')}", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Metadata
        metadata = doc.add_paragraph()
        metadata.add_run(f"Wygenerowano: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n").italic = True
        metadata.add_run("Platform: Sight AI Focus Groups").italic = True

        # Pytania i odpowiedzi
        doc.add_heading("Pytania i Odpowiedzi", level=2)

        for i, question in enumerate(survey_data.get("questions", []), 1):
            # Pytanie
            doc.add_heading(f"Pytanie {i}: {question.get('text', 'N/A')}", level=3)

            # Odpowiedzi
            if survey_data.get("responses"):
                for response in survey_data["responses"]:
                    if response.get("question_id") == question.get("id"):
                        p = doc.add_paragraph(style="List Bullet")
                        p.add_run(f"{response.get('persona_name', 'Unknown')}: ").bold = True
                        p.add_run(response.get("answer", ""))

        # Konwertuj do bytes
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_bytes = docx_buffer.getvalue()

        logger.info(f"DOCX wygenerowany, rozmiar: {len(docx_bytes)} bytes")
        return docx_bytes

    async def generate_project_docx(
        self,
        project_data: Dict[str, Any],
        user_tier: str = "free",
        include_full_personas: bool = False,
    ) -> bytes:
        """
        Generuje kompletny raport DOCX dla projektu.

        Args:
            project_data: Dane projektu (dict z polami: name, description, personas, focus_groups, surveys)
            user_tier: Tier użytkownika
            include_full_personas: Czy dołączyć pełne profile person (False = skrócone)

        Returns:
            bytes: Zawartość pliku DOCX
        """
        logger.info(f"Generowanie DOCX projektu: {project_data.get('name', 'Unknown')}")

        doc = Document()

        if user_tier == "free":
            self._add_watermark(doc, "SIGHT FREE TIER")

        # Tytuł
        title = doc.add_heading(f"Raport Projektu: {project_data.get('name', 'Unknown')}", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Metadata
        metadata = doc.add_paragraph()
        metadata.add_run(f"Wygenerowano: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n").italic = True
        metadata.add_run("Platform: Sight AI Focus Groups").italic = True

        # Opis projektu
        if project_data.get("description"):
            doc.add_heading("Opis Projektu", level=2)
            doc.add_paragraph(project_data["description"])

        # Grupa docelowa
        if project_data.get("target_audience"):
            doc.add_heading("Grupa Docelowa", level=2)
            doc.add_paragraph(project_data["target_audience"])

        # Cele badawcze
        if project_data.get("research_objectives"):
            doc.add_heading("Cele Badawcze", level=2)
            doc.add_paragraph(project_data["research_objectives"])

        # Agregacje i statystyki (używamy tych samych metod co PDF)
        personas = project_data.get("personas", [])
        focus_groups = project_data.get("focus_groups", [])
        surveys = project_data.get("surveys", [])

        demographics_stats = self._aggregate_demographics(personas)
        fg_insights = self._extract_focus_group_insights(focus_groups)
        survey_aggregates = self._aggregate_survey_responses(surveys)

        # Statystyki demograficzne
        doc.add_heading("Statystyki Demograficzne", level=2)
        table = doc.add_table(rows=4, cols=2)
        table.style = "Light Grid Accent 1"

        table.rows[0].cells[0].text = "Liczba person"
        table.rows[0].cells[1].text = str(demographics_stats.get("total_personas", 0))
        table.rows[1].cells[0].text = "Zakres wiekowy"
        table.rows[1].cells[1].text = demographics_stats.get("age_range", "N/A")
        table.rows[2].cells[0].text = "Średni wiek"
        table.rows[2].cells[1].text = str(demographics_stats.get("avg_age", "N/A"))
        table.rows[3].cells[0].text = "Walidacja statystyczna"
        table.rows[3].cells[1].text = "TAK" if project_data.get("is_statistically_valid") else "NIE"

        # Rozkład płci
        if demographics_stats.get("gender_distribution"):
            doc.add_heading("Rozkład Płci", level=3)
            gender_table = doc.add_table(rows=len(demographics_stats["gender_distribution"]) + 1, cols=2)
            gender_table.style = "Light Grid Accent 1"
            gender_table.rows[0].cells[0].text = "Płeć"
            gender_table.rows[0].cells[1].text = "Liczba"

            for i, (gender, count) in enumerate(demographics_stats["gender_distribution"].items(), 1):
                gender_table.rows[i].cells[0].text = str(gender)
                gender_table.rows[i].cells[1].text = str(count)

        # Rozkład wykształcenia
        if demographics_stats.get("education_distribution"):
            doc.add_heading("Rozkład Wykształcenia", level=3)
            edu_table = doc.add_table(rows=len(demographics_stats["education_distribution"]) + 1, cols=2)
            edu_table.style = "Light Grid Accent 1"
            edu_table.rows[0].cells[0].text = "Wykształcenie"
            edu_table.rows[0].cells[1].text = "Liczba"

            for i, (edu, count) in enumerate(demographics_stats["education_distribution"].items(), 1):
                edu_table.rows[i].cells[0].text = str(edu)
                edu_table.rows[i].cells[1].text = str(count)

        # Przykładowe persony
        personas_sample = personas[:10] if not include_full_personas else personas
        if personas_sample:
            heading_text = "Przykładowe Persony" if not include_full_personas else "Wszystkie Persony"
            if not include_full_personas:
                heading_text += " (Top 10)"
            doc.add_heading(heading_text, level=2)

            for persona in personas_sample:
                doc.add_heading(persona.get("full_name") or persona.get("name", "Unknown"), level=3)
                p = doc.add_paragraph()
                p.add_run("Wiek: ").bold = True
                p.add_run(f"{persona.get('age', 'N/A')}, ")
                p.add_run("Zawód: ").bold = True
                p.add_run(str(persona.get("occupation", "N/A")))

                p = doc.add_paragraph()
                p.add_run("Wykształcenie: ").bold = True
                p.add_run(f"{persona.get('education_level', 'N/A')}, ")
                p.add_run("Lokalizacja: ").bold = True
                p.add_run(str(persona.get("location", "N/A")))

                if persona.get("values"):
                    p = doc.add_paragraph()
                    p.add_run("Wartości: ").bold = True
                    p.add_run(", ".join(persona["values"]))

                if persona.get("interests"):
                    p = doc.add_paragraph()
                    p.add_run("Zainteresowania: ").bold = True
                    p.add_run(", ".join(persona["interests"]))

        # Grupy fokusowe
        if fg_insights.get("focus_groups_count", 0) > 0:
            doc.add_heading("Grupy Fokusowe - Kluczowe Insighty", level=2)

            p = doc.add_paragraph()
            p.add_run("Liczba grup: ").bold = True
            p.add_run(str(fg_insights["focus_groups_count"]))

            p = doc.add_paragraph()
            p.add_run("Łączna liczba uczestników: ").bold = True
            p.add_run(str(fg_insights["participant_count"]))

            # Główne tematy
            if fg_insights.get("key_themes"):
                doc.add_heading("Główne Tematy", level=3)
                for theme in fg_insights["key_themes"]:
                    doc.add_paragraph(theme, style="List Bullet")

            # Szczegóły grup
            for fg in focus_groups:
                doc.add_heading(fg.get("name", "Unknown"), level=3)

                p = doc.add_paragraph()
                p.add_run("Liczba uczestników: ").bold = True
                p.add_run(str(len(fg.get("persona_ids", []))))

                p = doc.add_paragraph()
                p.add_run("Liczba pytań: ").bold = True
                p.add_run(str(len(fg.get("questions", []))))

                if fg.get("ai_summary") and isinstance(fg["ai_summary"], dict):
                    if fg["ai_summary"].get("main_insights"):
                        p = doc.add_paragraph()
                        p.add_run("Główne wnioski: ").bold = True
                        p.add_run(fg["ai_summary"]["main_insights"])

        # Ankiety
        if survey_aggregates.get("surveys_count", 0) > 0:
            doc.add_heading("Ankiety - Podsumowanie", level=2)

            table = doc.add_table(rows=4, cols=2)
            table.style = "Light Grid Accent 1"
            table.rows[0].cells[0].text = "Liczba ankiet"
            table.rows[0].cells[1].text = str(survey_aggregates["surveys_count"])
            table.rows[1].cells[0].text = "Ankiety ukończone"
            table.rows[1].cells[1].text = str(survey_aggregates["completed_surveys"])
            table.rows[2].cells[0].text = "Łączna liczba odpowiedzi"
            table.rows[2].cells[1].text = str(survey_aggregates["total_responses"])

            # Szczegóły ankiet
            for survey in surveys:
                doc.add_heading(survey.get("title", "Unknown"), level=3)

                p = doc.add_paragraph()
                p.add_run("Status: ").bold = True
                p.add_run(survey.get("status", "N/A"))

                p = doc.add_paragraph()
                p.add_run("Liczba odpowiedzi: ").bold = True
                p.add_run(f"{survey.get('actual_responses', 0)} / {survey.get('target_responses', 0)}")

        # Podsumowanie
        doc.add_heading("Podsumowanie", level=2)
        summary_text = (
            f"Raport zawiera dane z {demographics_stats.get('total_personas', 0)} person, "
            f"{fg_insights.get('focus_groups_count', 0)} grup fokusowych "
            f"oraz {survey_aggregates.get('surveys_count', 0)} ankiet."
        )
        if project_data.get("is_statistically_valid"):
            summary_text += " Rozkład demograficzny został zwalidowany statystycznie (chi-kwadrat, p > 0.05)."
        doc.add_paragraph(summary_text)

        # Konwertuj do bytes
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_bytes = docx_buffer.getvalue()

        logger.info(f"DOCX projektu wygenerowany, rozmiar: {len(docx_bytes)} bytes")
        return docx_bytes

    def _aggregate_demographics(self, personas: list) -> Dict[str, Any]:
        """Agreguje statystyki demograficzne z listy person."""
        if not personas:
            return {}

        from collections import Counter

        ages = [p.get("age") for p in personas if p.get("age")]
        genders = [p.get("gender") for p in personas if p.get("gender")]
        educations = [p.get("education_level") for p in personas if p.get("education_level")]
        locations = [p.get("location") for p in personas if p.get("location")]

        return {
            "age_range": f"{min(ages)}-{max(ages)}" if ages else "N/A",
            "avg_age": round(sum(ages) / len(ages), 1) if ages else 0,
            "gender_distribution": dict(Counter(genders)),
            "education_distribution": dict(Counter(educations)),
            "top_locations": dict(Counter(locations).most_common(5)),
            "total_personas": len(personas),
        }

    def _extract_focus_group_insights(self, focus_groups: list) -> Dict[str, Any]:
        """Wyciąga kluczowe insighty z grup fokusowych."""
        if not focus_groups:
            return {"key_themes": [], "participant_count": 0}

        all_themes = []
        total_participants = 0

        for fg in focus_groups:
            if fg.get("ai_summary") and isinstance(fg["ai_summary"], dict):
                themes = fg["ai_summary"].get("key_themes", [])
                all_themes.extend(themes)
            total_participants += len(fg.get("persona_ids", []))

        return {
            "key_themes": list(set(all_themes))[:10],  # Top 10 unikalnych tematów
            "participant_count": total_participants,
            "focus_groups_count": len(focus_groups),
        }

    def _aggregate_survey_responses(self, surveys: list) -> Dict[str, Any]:
        """Agreguje odpowiedzi z ankiet."""
        if not surveys:
            return {"total_responses": 0, "surveys_count": 0}

        total_responses = sum(s.get("actual_responses", 0) for s in surveys)
        completed_surveys = sum(1 for s in surveys if s.get("status") == "completed")

        return {
            "total_responses": total_responses,
            "surveys_count": len(surveys),
            "completed_surveys": completed_surveys,
        }

    def _add_watermark(self, doc: Document, text: str):
        """
        Dodaje watermark do dokumentu DOCX.

        Note: python-docx nie ma natywnej obsługi watermarks,
        więc dodajemy jako header z szarym tekstem.
        """
        section = doc.sections[0]
        header = section.header
        watermark_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        watermark_para.text = text
        watermark_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Stylizacja watermark
        run = watermark_para.runs[0]
        run.font.size = Pt(48)
        run.font.color.rgb = RGBColor(200, 200, 200)  # Szary
        run.font.bold = True
