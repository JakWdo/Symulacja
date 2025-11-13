"""
Testy jednostkowe dla DOCXGenerator

Testuje generowanie raportów DOCX (Microsoft Word) dla projektów.
"""

import pytest
from io import BytesIO
from docx import Document
from app.services.export.docx_generator import DOCXGenerator


class TestDOCXGenerator:
    """Testy dla generatora DOCX."""

    @pytest.fixture
    def docx_generator(self):
        """Fixture zwracający instancję DOCXGenerator."""
        return DOCXGenerator()

    @pytest.fixture
    def sample_project_data(self):
        """Fixture z przykładowymi danymi projektu."""
        return {
            "id": "test-project-id",
            "name": "Testowy Projekt DOCX",
            "description": "Opis testowego projektu",
            "target_audience": "Młodzi profesjonaliści",
            "research_objectives": "Zrozumienie potrzeb użytkowników",
            "is_statistically_valid": True,
            "personas": [
                {
                    "id": "persona-1",
                    "full_name": "Jan Kowalski",
                    "name": "Jan Kowalski",
                    "age": 28,
                    "gender": "male",
                    "occupation": "Software Developer",
                    "education_level": "Bachelor's Degree",
                    "location": "Warsaw",
                    "values": ["Innovation", "Growth"],
                    "interests": ["Technology", "Travel"],
                },
            ],
            "focus_groups": [
                {
                    "id": "fg-1",
                    "name": "Grupa Fokusowa 1",
                    "persona_ids": ["persona-1"],
                    "questions": ["Pytanie 1?"],
                    "ai_summary": {
                        "main_insights": "Kluczowe wnioski",
                        "key_themes": ["Temat 1"],
                    },
                }
            ],
            "surveys": [
                {
                    "id": "survey-1",
                    "title": "Ankieta testowa",
                    "status": "completed",
                    "actual_responses": 50,
                    "target_responses": 100,
                }
            ],
        }

    @pytest.mark.asyncio
    async def test_generate_project_docx_returns_bytes(
        self, docx_generator, sample_project_data
    ):
        """Test generowania DOCX - zwraca niepusty plik."""
        docx_bytes = await docx_generator.generate_project_docx(
            project_data=sample_project_data,
            user_tier="free",
        )

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

    @pytest.mark.asyncio
    async def test_generate_project_docx_is_valid_document(
        self, docx_generator, sample_project_data
    ):
        """Test czy wygenerowany DOCX jest poprawnym dokumentem Word."""
        docx_bytes = await docx_generator.generate_project_docx(
            project_data=sample_project_data,
            user_tier="free",
        )

        # Spróbuj otworzyć jako dokument Word
        doc = Document(BytesIO(docx_bytes))

        # Dokument powinien mieć paragrafy
        assert len(doc.paragraphs) > 0

    @pytest.mark.asyncio
    async def test_generate_project_docx_contains_project_name(
        self, docx_generator, sample_project_data
    ):
        """Test czy DOCX zawiera nazwę projektu."""
        docx_bytes = await docx_generator.generate_project_docx(
            project_data=sample_project_data,
            user_tier="free",
        )

        doc = Document(BytesIO(docx_bytes))

        # Nazwa projektu powinna być w którymś paragrafie
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Testowy Projekt DOCX" in all_text

    @pytest.mark.asyncio
    async def test_generate_project_docx_contains_sections(
        self, docx_generator, sample_project_data
    ):
        """Test czy DOCX zawiera wymagane sekcje."""
        docx_bytes = await docx_generator.generate_project_docx(
            project_data=sample_project_data,
            user_tier="free",
        )

        doc = Document(BytesIO(docx_bytes))
        all_text = "\n".join([p.text for p in doc.paragraphs])

        # Sprawdź kluczowe sekcje
        assert "Opis Projektu" in all_text
        assert "Statystyki Demograficzne" in all_text
        assert "Przykładowe Persony" in all_text or "Wszystkie Persony" in all_text
        assert "Podsumowanie" in all_text

    @pytest.mark.asyncio
    async def test_generate_project_docx_with_full_personas(
        self, docx_generator, sample_project_data
    ):
        """Test generowania DOCX z pełnymi personami."""
        docx_bytes = await docx_generator.generate_project_docx(
            project_data=sample_project_data,
            user_tier="pro",
            include_full_personas=True,
        )

        doc = Document(BytesIO(docx_bytes))
        all_text = "\n".join([p.text for p in doc.paragraphs])

        assert "Wszystkie Persony" in all_text or "Jan Kowalski" in all_text

    @pytest.mark.asyncio
    async def test_generate_project_docx_empty_project(self, docx_generator):
        """Test generowania DOCX dla projektu bez danych."""
        empty_project = {
            "id": "empty-id",
            "name": "Pusty Projekt",
            "description": "",
            "personas": [],
            "focus_groups": [],
            "surveys": [],
        }

        docx_bytes = await docx_generator.generate_project_docx(
            project_data=empty_project,
            user_tier="free",
        )

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

        # Sprawdź czy można otworzyć dokument
        doc = Document(BytesIO(docx_bytes))
        assert len(doc.paragraphs) > 0

    def test_aggregate_demographics(self, docx_generator):
        """Test agregacji statystyk demograficznych."""
        personas = [
            {"age": 25, "gender": "male", "education_level": "Bachelor's", "location": "Warsaw"},
            {"age": 35, "gender": "female", "education_level": "PhD", "location": "Gdansk"},
        ]

        stats = docx_generator._aggregate_demographics(personas)

        assert stats["total_personas"] == 2
        assert stats["age_range"] == "25-35"
        assert stats["avg_age"] == 30.0

    def test_extract_focus_group_insights(self, docx_generator):
        """Test ekstrakcji insightów z grup fokusowych."""
        focus_groups = [
            {
                "persona_ids": ["p1", "p2", "p3"],
                "ai_summary": {"key_themes": ["Innovation", "Growth"]},
            },
        ]

        insights = docx_generator._extract_focus_group_insights(focus_groups)

        assert insights["focus_groups_count"] == 1
        assert insights["participant_count"] == 3

    def test_aggregate_survey_responses(self, docx_generator):
        """Test agregacji odpowiedzi z ankiet."""
        surveys = [
            {"status": "completed", "actual_responses": 100},
            {"status": "running", "actual_responses": 50},
        ]

        aggregates = docx_generator._aggregate_survey_responses(surveys)

        assert aggregates["surveys_count"] == 2
        assert aggregates["total_responses"] == 150
        assert aggregates["completed_surveys"] == 1
